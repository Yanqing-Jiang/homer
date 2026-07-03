#!/usr/bin/env python3
"""
Build an ATS-safe Calibri 11pt reference.docx for Pandoc to use as --reference-doc.

What Pandoc does with this file:
- Reads style definitions (Normal, Heading 1, Heading 2, Heading 3, List Paragraph).
- Reads page margins.
- Ignores all content. Only styles propagate.

ATS-safety choices:
- Calibri 11pt body, 12pt bold section headings, 22pt bold name.
- 0.65" margins all around.
- Single column, no headers, no footers, no tables.
- Black text only, no fancy colors.
- Spacing tuned for 2-page senior resumes.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from pathlib import Path

OUT = str(Path(__file__).resolve().parents[1] / "templates" / "reference.docx")
FONT = "Calibri"
BLACK = RGBColor(0, 0, 0)


def set_font(run, name, size_pt, bold=False, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color


def configure_style(styles, name, *, font_size, bold=False, space_before=0, space_after=0,
                    keep_with_next=False):
    """Configure an existing style with Calibri + size + spacing."""
    style = styles[name]
    style.font.name = FONT
    style.font.size = Pt(font_size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if keep_with_next:
        pf.keep_with_next = True


def main():
    doc = Document()

    # --- Page setup: 0.65" margins ---
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # --- Configure built-in styles ---
    styles = doc.styles

    # Body text
    configure_style(styles, "Normal", font_size=11, space_after=4)

    # Heading 1 → resume header (name)
    configure_style(styles, "Heading 1", font_size=22, bold=True,
                    space_before=0, space_after=2, keep_with_next=True)

    # Heading 2 → section names (Summary, Skills, Experience, ...)
    configure_style(styles, "Heading 2", font_size=12, bold=True,
                    space_before=10, space_after=3, keep_with_next=True)

    # Heading 3 → job entry headers (employer + role)
    configure_style(styles, "Heading 3", font_size=11, bold=True,
                    space_before=6, space_after=1, keep_with_next=True)

    # List Paragraph → bullets
    configure_style(styles, "List Paragraph", font_size=11,
                    space_before=0, space_after=2)

    # --- Sample content (Pandoc IGNORES this; it's just for visual inspection) ---
    # If you open this file in Word, you'll see what each style renders like.

    h1 = doc.add_heading("Yanqing Jiang", level=1)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    contact = doc.add_paragraph(
        "Seattle, WA | jiangyanqing90@gmail.com | (270) 978-9240 | "
        "yanqing.app | https://www.linkedin.com/in/jiangyanqing/"
    )
    contact.paragraph_format.space_after = Pt(2)

    headline = doc.add_paragraph()
    r = headline.add_run("Sample Headline | Style Reference Document")
    r.bold = True
    set_font(r, FONT, 11, bold=True)
    headline.paragraph_format.space_after = Pt(8)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "This file's content is sample text for visual inspection only. "
        "Pandoc reads ONLY the styles from this file and ignores all content. "
        "When the resume-tailor skill runs, it produces tailored markdown, "
        "then renders to .docx using this reference for fonts, sizes, margins, "
        "and spacing."
    )

    doc.add_heading("Skills", level=2)
    skills_para = doc.add_paragraph()
    r1 = skills_para.add_run("Analytics and AI: ")
    r1.bold = True
    skills_para.add_run("A/B testing, hypothesis testing, machine learning pipelines")

    doc.add_heading("Experience", level=2)
    doc.add_heading("Procter & Gamble — Advanced Analytics Senior Manager", level=3)
    sub = doc.add_paragraph("Amazon Sales Team | Seattle, WA | Jul 2024 – Present")
    sub.paragraph_format.space_after = Pt(3)
    bullet1 = doc.add_paragraph(
        "Sample bullet using the List Paragraph style — this is the style that "
        "Pandoc will apply to all dash-prefixed bullets in the markdown source.",
        style="List Paragraph",
    )
    bullet1.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading("Education", level=2)
    edu = doc.add_paragraph(
        "The Ohio State University, Max M. Fisher College of Business — "
        "Bachelor of Science, Corporate Finance (Jan 2012 – Dec 2013)",
        style="List Paragraph",
    )
    edu.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading("Certifications", level=2)
    for cert in [
        "Machine Learning Specialization — DeepLearning.AI",
        "Data Analyst Professional Certificate — IBM",
    ]:
        p = doc.add_paragraph(cert, style="List Paragraph")
        p.paragraph_format.left_indent = Inches(0.25)

    # --- Save ---
    doc.save(OUT)
    print(f"✅ Wrote {OUT}")


if __name__ == "__main__":
    main()
